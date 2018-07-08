#!coding:utf-8
import json
import os

import numpy as np
import matplotlib.pyplot as plt
import matplotlib as mpl
import matplotlib.font_manager
from docx import Document
from docx.shared import Pt
from docx.shared import Inches
from docx.oxml.ns import qn

from numpy.ctypeslib import load_library, ndpointer

from getpredata import get_json_file_path, extract_pressure, get_dir_names

mpl.rcParams['font.size'] = 15
custom_font = mpl.font_manager.FontProperties(fname='/home/ri/user/Tem/微软雅黑+Arial.ttf')


def get_info_from_pa_path(pa_path):
    """
    获取患者信息通过文件路径
    :param pa_path: 文件路径
    :return:    patient_name[1] 患者姓名
                patient_condition[1] 患者状态
                patient_fileid[1] 文件编号
    """
    patient_fileid = os.path.split(os.path.splitext(pa_path)[0])
    patient_name = os.path.split(patient_fileid[0])
    patient_condition = os.path.split(patient_name[0])
    return patient_name[1], patient_condition, patient_fileid[1]


def get_lib_interpolate_foot(core_type='cubic'):
    """ 读取C文件
        arg[8],empty_arg[660]
    """
    lib = load_library('/home/ri/cubic.so', '.')
    if core_type == 'line':
        lib = load_library('/home/ri/line.so', '.')
    interpolate_foot = lib.interpolate_foot
    interpolate_foot.argtypes = [ndpointer(dtype='f4', ndim=1, flags='C_CONTIGUOUS'),
                                 ndpointer(dtype='f4', ndim=1, flags='C_CONTIGUOUS')]
    interpolate_foot.restype = None
    return interpolate_foot


def print_data_format(data):
    return "HL:%(d0)g \nHM:%(d1)g \nM4,5:%(d2)g \nM2,3:%(d3)g\nM1:%(d4)g \nMID:%(d5)g,%(d6)g \nT1:%(d7)g " \
           % {'d0': data[0], 'd1': data[1], 'd2': data[2], 'd3': data[3],
              'd4': data[4], 'd5': data[5], 'd6': data[6], 'd7': data[7]}


def save_2d_pressure(pa_data, pa_info, lib, document):
    grid = np.empty(660, dtype='f4')
    lib(np.array(pa_data).astype('f4'), grid)
    plt.title(pa_info, fontproperties=custom_font)
    plt.imshow(grid.reshape((44, 15)), interpolation='nearest', origin='lower', cmap=plt.get_cmap('jet'))
    plt.colorbar()
    plt.text(-13, 0, print_data_format(pa_data), horizontalalignment='left', fontsize=10, fontproperties=custom_font)
    plt.savefig('/home/ri/a.png')
    # plt.show()
    plt.clf()
    document.add_picture('/home/ri/a.png', width=Inches(5.5))


def make_2d_pressure(path):
    before_set = set()
    after_set = set()
    novel_data = get_ordinary_data(novel_path)
    document = Document()
    jsons_path = get_json_file_path(path)
    lib_if = get_lib_interpolate_foot()
    for json_path in jsons_path:
        print(json_path)
        pa_data = extract_pressure(json_path.replace('.json', '.dat'))
        with open(json_path, 'r') as load_file:
            load_json = json.load(load_file)
            pa_data_step = load_json['footstep']
        pa_name, pa_condition, pa_fileid = get_info_from_pa_path(json_path)
        for i in range(len(pa_data_step)):
            pa_info = 'Condition :' + str(pa_condition[1]) \
                      + ' Name :' + str(pa_name) \
                      + ' Step :' + str((i + 1)) \
                      + '\n File :' + str(pa_fileid)
            step_begin = pa_data_step[i][0]
            step_end = pa_data_step[i][1]
            pa_step_data = [max(pa_data[i][step_begin:step_end]) for i in range(8)]
            width = 3.5
            info = 'Condition :' + str(pa_condition[1]) + ' Name :' + str(pa_name)
            if pa_condition[1] == '术前':
                b_s_size = len(before_set)
                before_set.add(pa_name)
                if len(before_set) != b_s_size:
                    if pa_name in novel_data:
                        p_path = os.path.join(novel_data[pa_name], pa_name)
                        left_path = os.path.join(p_path, 'left.jpg')
                        right_path = os.path.join(p_path, 'right.jpg')
                        document.add_paragraph(info + ' left')
                        document.add_picture(left_path, width=Inches(width))
                        document.add_page_break()
                        document.add_paragraph(info + ' right')
                        document.add_picture(right_path, width=Inches(width))
            elif pa_condition[1] == '术后':
                t_name = pa_name + '1'
                a_s_size = len(after_set)
                after_set.add(t_name)
                if len(after_set) != a_s_size:
                    if t_name in novel_data:
                        p_path = os.path.join(novel_data[pa_name], t_name)
                        left_path = os.path.join(p_path, 'left.jpg')
                        right_path = os.path.join(p_path, 'right.jpg')
                        document.add_paragraph(info + ' left')
                        document.add_picture(left_path, width=Inches(width))
                        document.add_page_break()
                        document.add_paragraph(info + ' right')
                        document.add_picture(right_path, width=Inches(width))
            save_2d_pressure(pa_step_data, pa_info, lib_if, document)
    document.save('/home/ri/demo.docx')


def get_ordinary_data(path):
    day_name = {}
    day_path = get_dir_names(path)
    for day in day_path:
        a_day = os.path.join(path, day)
        name_path = get_dir_names(a_day)
        for name in name_path:
            day_name[name] = a_day
    return day_name


if __name__ == '__main__':
    patient_path = '/home/ri/Data/patient/足根'
    novel_path = '/home/ri/user/Tem/医院数据/积水潭/novel'
    make_2d_pressure(patient_path)
    print('完成')
